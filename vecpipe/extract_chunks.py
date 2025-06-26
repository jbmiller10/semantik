#!/usr/bin/env python3
"""
Document extraction and chunking module V2
Uses tiktoken for accurate token counting
Implements file change tracking with SHA256
"""

import os
import hashlib
import logging
import json
from pathlib import Path
from typing import List, Dict, Optional, Generator, Tuple
from datetime import datetime
import pyarrow as pa
import pyarrow.parquet as pq
from tqdm import tqdm
import argparse
import tiktoken

# Document parsing libraries
from pypdf import PdfReader
from docx import Document as DocxDocument

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Constants
DEFAULT_CHUNK_SIZE = 600  # tokens
DEFAULT_CHUNK_OVERLAP = 200  # tokens
OUTPUT_DIR = "/opt/vecpipe/extract"
ERROR_LOG = "/tmp/error_extract.log"
FILE_TRACKING_DB = "/var/embeddings/file_tracking.json"

class TokenChunker:
    """Chunk text by token count using tiktoken"""
    
    def __init__(self, model_name: str = "cl100k_base", chunk_size: int = 600, chunk_overlap: int = 200):
        """Initialize tokenizer for chunking"""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Use tiktoken for accurate token counting
        try:
            self.tokenizer = tiktoken.get_encoding(model_name)
        except:
            # Fallback to default encoding
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        logger.info(f"Initialized tokenizer: {model_name}, chunk_size: {chunk_size}, overlap: {chunk_overlap}")
    
    def chunk_text(self, text: str, doc_id: str) -> List[Dict]:
        """Split text into overlapping chunks by token count"""
        if not text.strip():
            return []
        
        # Tokenize entire text
        tokens = self.tokenizer.encode(text)
        
        if len(tokens) <= self.chunk_size:
            # Text fits in single chunk
            return [{
                'doc_id': doc_id,
                'chunk_id': f"{doc_id}_0000",
                'text': text.strip(),
                'token_count': len(tokens),
                'start_token': 0,
                'end_token': len(tokens)
            }]
        
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(tokens):
            # Determine chunk boundaries
            end = min(start + self.chunk_size, len(tokens))
            
            # Extract chunk tokens
            chunk_tokens = tokens[start:end]
            
            # Decode back to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            # Try to break at sentence boundary if not at end
            if end < len(tokens):
                # Look for sentence end in last 10% of chunk
                search_start = int(len(chunk_tokens) * 0.9)
                best_break = len(chunk_tokens)
                
                for i in range(search_start, len(chunk_tokens)):
                    decoded = self.tokenizer.decode(chunk_tokens[:i])
                    if decoded.rstrip().endswith(('.', '!', '?', '\n\n')):
                        best_break = i
                        break
                
                # Adjust chunk if we found a better break point
                if best_break < len(chunk_tokens):
                    chunk_tokens = chunk_tokens[:best_break]
                    chunk_text = self.tokenizer.decode(chunk_tokens)
                    end = start + best_break
            
            chunks.append({
                'doc_id': doc_id,
                'chunk_id': f"{doc_id}_{chunk_id:04d}",
                'text': chunk_text.strip(),
                'token_count': len(chunk_tokens),
                'start_token': start,
                'end_token': end
            })
            
            chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            
            # Ensure progress
            if start <= end - self.chunk_size:
                start = end
        
        logger.debug(f"Created {len(chunks)} chunks from {len(tokens)} tokens")
        return chunks

class FileChangeTracker:
    """Track file changes using SHA256 and SCD-like approach"""
    
    def __init__(self, db_path: str = FILE_TRACKING_DB):
        self.db_path = db_path
        self.tracking_data = self._load_tracking_data()
    
    def _load_tracking_data(self) -> Dict:
        """Load tracking data from JSON file"""
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, 'r') as f:
                    return json.load(f)
            except:
                logger.warning(f"Failed to load tracking data from {self.db_path}")
        return {"files": {}}
    
    def _save_tracking_data(self):
        """Save tracking data to JSON file"""
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        with open(self.db_path, 'w') as f:
            json.dump(self.tracking_data, f, indent=2)
    
    def get_file_hash(self, filepath: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(filepath, "rb") as f:
            for byte_block in iter(lambda: f.read(65536), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def should_process_file(self, filepath: str) -> Tuple[bool, Optional[str]]:
        """Check if file should be processed based on hash"""
        current_hash = self.get_file_hash(filepath)
        file_key = str(Path(filepath).absolute())
        
        if file_key in self.tracking_data["files"]:
            file_info = self.tracking_data["files"][file_key]
            if file_info["hash"] == current_hash:
                # File unchanged
                return False, current_hash
            else:
                # File changed
                logger.info(f"File changed: {filepath}")
                return True, current_hash
        else:
            # New file
            logger.info(f"New file: {filepath}")
            return True, current_hash
    
    def update_file_tracking(self, filepath: str, file_hash: str, doc_id: str, chunks_created: int):
        """Update tracking information for a file"""
        file_key = str(Path(filepath).absolute())
        now = datetime.now().isoformat()
        
        if file_key in self.tracking_data["files"]:
            # Update existing entry
            file_info = self.tracking_data["files"][file_key]
            file_info["hash"] = file_hash
            file_info["last_seen"] = now
            file_info["last_processed"] = now
            file_info["doc_id"] = doc_id
            file_info["chunks_created"] = chunks_created
            file_info["process_count"] = file_info.get("process_count", 0) + 1
        else:
            # Create new entry
            self.tracking_data["files"][file_key] = {
                "hash": file_hash,
                "first_seen": now,
                "last_seen": now,
                "last_processed": now,
                "doc_id": doc_id,
                "chunks_created": chunks_created,
                "process_count": 1
            }
        
        self._save_tracking_data()
    
    def get_removed_files(self, current_files: List[str]) -> List[Dict]:
        """Find files that were tracked but no longer exist"""
        current_file_keys = {str(Path(f).absolute()) for f in current_files}
        removed_files = []
        
        for file_key, file_info in self.tracking_data["files"].items():
            if file_key not in current_file_keys:
                removed_files.append({
                    "path": file_key,
                    "doc_id": file_info.get("doc_id"),
                    "last_seen": file_info.get("last_seen")
                })
        
        return removed_files
    
    def remove_file(self, filepath: str):
        """Remove a file from tracking"""
        file_key = str(Path(filepath).absolute()) if not os.path.isabs(filepath) else filepath
        if file_key in self.tracking_data["files"]:
            del self.tracking_data["files"][file_key]
            logger.info(f"Removed file from tracking: {file_key}")
    
    def save(self):
        """Save tracking data to disk"""
        self._save_tracking_data()

def extract_text_pdf(filepath: str) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(filepath)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num} from {filepath}: {e}")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract PDF {filepath}: {e}")
        raise

def extract_text_docx(filepath: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(filepath)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to extract DOCX {filepath}: {e}")
        raise

def extract_text_txt(filepath: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to extract TXT {filepath}: {e}")
        raise

def extract_text(filepath: str) -> str:
    """Extract text from file based on extension"""
    ext = Path(filepath).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return extract_text_docx(filepath)
    elif ext in ['.txt', '.text']:
        return extract_text_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def process_file_v2(filepath: str, output_dir: str, chunker: TokenChunker, tracker: FileChangeTracker) -> Optional[str]:
    """Process a single file with change tracking"""
    try:
        # Check if file needs processing
        should_process, file_hash = tracker.should_process_file(filepath)
        
        if not should_process:
            logger.info(f"Skipping unchanged file: {filepath}")
            return None
        
        # Generate document ID from file path
        doc_id = hashlib.md5(filepath.encode()).hexdigest()[:16]
        
        # Check if output already exists (for resume capability)
        output_path = os.path.join(output_dir, f"{doc_id}.parquet")
        if os.path.exists(output_path):
            # Verify the existing output is valid
            try:
                existing_table = pq.read_table(output_path)
                if len(existing_table) > 0:
                    logger.info(f"Valid output exists for: {filepath}")
                    tracker.update_file_tracking(filepath, file_hash, doc_id, len(existing_table))
                    return output_path
            except:
                logger.warning(f"Invalid existing output, reprocessing: {filepath}")
        
        # Extract text
        logger.info(f"Extracting: {filepath}")
        text = extract_text(filepath)
        
        if not text.strip():
            logger.warning(f"No text extracted from: {filepath}")
            tracker.update_file_tracking(filepath, file_hash, doc_id, 0)
            return None
        
        # Chunk text using token-based chunking
        chunks = chunker.chunk_text(text, doc_id)
        logger.info(f"Created {len(chunks)} chunks from: {filepath}")
        
        # Convert to parquet
        if chunks:
            df_data = {
                'doc_id': [c['doc_id'] for c in chunks],
                'chunk_id': [c['chunk_id'] for c in chunks],
                'path': [filepath] * len(chunks),
                'text': [c['text'] for c in chunks],
                'token_count': [c['token_count'] for c in chunks],
                'file_hash': [file_hash] * len(chunks)
            }
            
            table = pa.table(df_data)
            pq.write_table(table, output_path)
            
            # Update tracking
            tracker.update_file_tracking(filepath, file_hash, doc_id, len(chunks))
            
            return output_path
        
        return None
        
    except Exception as e:
        logger.error(f"Failed to process {filepath}: {e}")
        # Log error to file
        with open(ERROR_LOG, 'a') as f:
            f.write(f"{filepath}\t{str(e)}\n")
        return None

# Compatibility wrappers for old API
_default_chunker = None

def chunk_text(text: str, doc_id: str) -> List[Dict]:
    """Compatibility wrapper for old chunk_text function"""
    global _default_chunker
    if _default_chunker is None:
        _default_chunker = TokenChunker()
    return _default_chunker.chunk_text(text, doc_id)

def process_file(filepath: str, output_dir: str) -> Optional[str]:
    """Compatibility wrapper for old process_file function"""
    global _default_chunker
    if _default_chunker is None:
        _default_chunker = TokenChunker()
    tracker = FileChangeTracker()
    return process_file_v2(filepath, output_dir, _default_chunker, tracker)

def main():
    parser = argparse.ArgumentParser(description="Extract and chunk documents (V2)")
    parser.add_argument('--input', '-i', required=True, help='Input file list (null-delimited)')
    parser.add_argument('--output', '-o', default=OUTPUT_DIR, help='Output directory for parquet files')
    parser.add_argument('--chunk-size', '-cs', type=int, default=DEFAULT_CHUNK_SIZE, help='Chunk size in tokens')
    parser.add_argument('--chunk-overlap', '-co', type=int, default=DEFAULT_CHUNK_OVERLAP, help='Chunk overlap in tokens')
    parser.add_argument('--model-encoding', '-m', default='cl100k_base', help='Tiktoken model encoding')
    parser.add_argument('--batch-size', '-b', type=int, default=100, help='Files to process per batch')
    
    args = parser.parse_args()
    
    # Initialize components
    chunker = TokenChunker(
        model_name=args.model_encoding,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap
    )
    tracker = FileChangeTracker()
    
    # Ensure output directory exists
    os.makedirs(args.output, exist_ok=True)
    
    # Read file list
    with open(args.input, 'rb') as f:
        file_list = f.read().decode('utf-8').split('\0')
        file_list = [f for f in file_list if f]  # Remove empty strings
    
    logger.info(f"Found {len(file_list)} files to process")
    
    # Check for removed files
    removed_files = tracker.get_removed_files(file_list)
    if removed_files:
        logger.info(f"Found {len(removed_files)} removed files")
        for removed in removed_files:
            logger.info(f"File removed: {removed['path']} (doc_id: {removed['doc_id']})")
    
    # Process files
    processed_count = 0
    skipped_count = 0
    failed_count = 0
    
    for filepath in tqdm(file_list, desc="Processing files"):
        result = process_file_v2(filepath, args.output, chunker, tracker)
        if result:
            processed_count += 1
        elif result is None:
            failed_count += 1
        else:
            skipped_count += 1
    
    logger.info(f"Processing complete: {processed_count} processed, {skipped_count} skipped, {failed_count} failed")

if __name__ == "__main__":
    main()